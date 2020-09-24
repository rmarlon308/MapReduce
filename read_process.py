from exceptions import Error
import os
import docx
from retrying import retry
import time

class ReadDocument(Error):
	n_part = 0

	def __init__(self,document_name, lines = 25, bug = False, testMode = False):
		self.bug = bug
		self.lines = lines
		self.document_name = document_name
		self.testMode = testMode

	def start(self):
		types = self.document_name.split(".")

		if types[len(types) - 1] == "txt":
			self.fileDistributor()
		else:
			self.readWord()

	@retry
	def readWord(self):
		doc = docx.Document(self.document_name)

		current_line = 0
		self.n_part = 1

		new_file = self.newFile(str(self.n_part) + ".txt")

		file = open(self.document_name, "r")

		for paragraph in doc.paragraphs:

			self.ifTestMode()

			if self.bug:
				self.setBug(False)
				self.raiseExeption("Error: Error en la division de archivos")

			current_line += 1

			if current_line % self.lines == 0:
				new_file.close()
				self.n_part += 1
				new_file = self.newFile(str(self.n_part) + ".txt")

			new_file.write(self.removeNoWords(paragraph.text))

		new_file.close()
		print("Tarea 1: Lectura de documento y division en archivos completada")




	def newFile(self, file_name):
		if os.path.exists(file_name):
			os.remove(file_name)
		return(open(file_name, "a+"))

	def openFile(self, file_name):
		return(open(file_name, "r"))

	def removeNoWords(self, text):
		chars = "&#()[],.`'!?;-_*/$:”‘“—’" + '"' + ""
		for c in chars:
			text = text.replace(c,"")
		return(text)

	def ifTestMode(self):
		if self.testMode:
			time.sleep(.3)


	@retry
	def fileDistributor(self):

		current_line = 0
		self.n_part = 1

		new_file = self.newFile(str(self.n_part) + ".txt")

		file = open(self.document_name, "r")

		count = 0

		while True:

			self.ifTestMode()
	
			if self.bug:
				self.setBug(False)
				self.raiseExeption("Error: Error en la division de archivos")

			count += 1

			line = file.readline()

			if not line:
				new_file.close()
				break

			line = self.removeNoWords(line)

			if line == "\n":
				continue

			current_line += 1

			if current_line % self.lines == 0:
				new_file.close()
				self.n_part += 1
				new_file = self.newFile(str(self.n_part) + ".txt")

			if count == 2:
				current_line = 1

			new_file.write(line)


		print("Tarea 1: Lectura de documento y division en archivos completada")
		file.close()

	def getBug(self):
		return(self.bug)

	def setBug(self, bug):
		self.bug = bug

	def getNPart(self):
		return self.n_part

	def getTaskCompleted(self):
		return self.taskCompleted
		




